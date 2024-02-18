# pip install python-docx
# pip install pywin32
# pip install natsort
# pip install Pillow
# https://python-docx.readthedocs.io/en/latest/user/text.html
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.text.hyperlink import CT_Hyperlink
from docx.enum.style import WD_STYLE_TYPE

from pathlib import Path
from PIL import Image
import os
import win32com.client
import shutil
import sys
from natsort import ns, natsort_keygen
import re
import analizze
out_ = []
def add_hyperlink(paragraph, text, url):
    """
    Добавляет гиперссылку в параграф.
    """
    run = paragraph.add_run(text)
    # run.hyperlink = CT_Hyperlink
    hyperlink = run._element
    r = hyperlink
    rPr = r.get_or_add_rPr()
    hlinkClick = OxmlElement('w:hlinkClick')
    hlinkClick.set('{http://www.w3.org/1999/xlink}href', url)
    rPr.append(hlinkClick)
    run.font.color.rgb = RGBColor(13, 109, 231)
    run.add_break()
    return run

def image_to_jpg(image_path):
    path = Path(image_path)
    # if path.suffix not in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
    jpg_image_path = f'{path.parent / path.stem}.jpg'
    Image.open(image_path).convert('RGB').save(jpg_image_path)
    return jpg_image_path
    # return image_path

def ConvertRtfToDocx(rootDir, file):
    if os.path.exists(rootDir + "\\{f_}".format(f_= file.replace('doc', 'docx').replace('rtf', 'docx'))):
        if os.path.exists(rootDir + "\\" + file):
            shutil.move(rootDir + "\\" + file , rootDir + "\\arch\\" + file)
        return True
    word = win32com.client.Dispatch("Word.Application")
    wdFormatDocumentDefault = 16
    wdHeaderFooterPrimary = 1
    doc = word.Documents.Open(rootDir + "\\" + file)
    for pic in doc.InlineShapes:
        pic.LinkFormat.SavePictureWithDocument = True
    for hPic in doc.sections(1).headers(wdHeaderFooterPrimary).Range.InlineShapes:
        hPic.LinkFormat.SavePictureWithDocument = True
    doc.SaveAs(str(rootDir + "\\{f_}".format(f_= file.replace('doc', 'docx').replace('rtf', 'docx'))), FileFormat=wdFormatDocumentDefault)
    doc.Close()
    word.Quit()
    shutil.move(rootDir + "\\" + file , rootDir + "\\arch\\" + file)
    return True

def open_doc(file):
    doc = open(file,"rb")
    return Document(doc)
def parse_table(document):
    style = document.styles.add_style('ReplClass1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(12)
    style.font.name = 'Times New Roman'
    style.font.bold = True

    style2 = document.styles.add_style('ReplClass2', WD_STYLE_TYPE.PARAGRAPH)
    style2.font.size = Pt(12)
    style2.font.name = 'Times New Roman'
    
    style3 = document.styles.add_style('ReplClass3', WD_STYLE_TYPE.PARAGRAPH)
    style3 = style2
    
    for table in document.tables:
        count = 0
        row1, row2, row3 = False, False, False
        for rows in table.rows:
            _cell = ()
            for cell in rows.cells:
                if _cell == cell:
                    continue
                _cell = cell
                
                # cell.text += '__test__'
                count_p = 0
                for paragraph in cell.paragraphs:
                    count_p += 1
                    
                    if 'Місцезнаходження' in paragraph.text:
                        row1 = True
                    if 'Інформація щодо режиму роботи' in paragraph.text:
                        row2 = True
                    if 'Телефон / факс, електронна  адреса, офіційний веб-сайт' in paragraph.text or 'Телефон / факс, адреса електронної пошти та веб-сайт' in paragraph.text:
                        row3 = True
                    for link in paragraph.hyperlinks:
                        pass
                        
                     
                    if row1 and 'сільської ради' in paragraph.text:
                        tmp_format = style
                        paragraph.text = paragraph.text.replace('______ сільської ради', 'Успенівської сільської ради')
                        # paragraph.style.element = tmp_format.element
                        paragraph.style = tmp_format
                        row1 = False
                        # print([paragraph.text, [paragraph.style.font.size]])
                    elif row2 and 'сільської ради' in paragraph.text:
                        tmp_format = style2
                        paragraph.text = paragraph.text.replace('______ сільської ради', 'Успенівської сільської ради')
                        # paragraph.style.element = tmp_format.element
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.style = style
                        if len(cell.paragraphs) > count_p:
                            paragraph2 = cell.paragraphs[count_p]
                            # print(count_p)
                        else:
                            paragraph2 = cell.add_paragraph()
                        paragraph2.text = ''
                        paragraph2.paragraph_format.space_before = Pt(0)
                        paragraph2.add_run()
                        run = paragraph2.runs[0]
                        run.add_text('Понеділок-П’ятниця: 8:00-17:00')
                        run.add_break()
                        run.add_text('Обідня перерва: 13:00-14:00')
                        paragraph2.style = tmp_format
                        paragraph2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        row2 = False
                        # print([paragraph.text, [paragraph.style.font.size]])
                    elif row3 and 'сільської ради' in paragraph.text:
                        tmp_format = style3
                        paragraph.text = paragraph.text.replace('______ сільської ради', 'Успенівської сільської ради')
                        # paragraph.style.element = tmp_format.element
                        paragraph.style = style
                        if len(cell.paragraphs) > count_p:
                            paragraph3 = cell.paragraphs[count_p]
                        else:
                            paragraph3 = cell.add_paragraph()
                        paragraph3.text = ''
                        paragraph3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = paragraph3.add_run()
                        # print(dir(paragraph3))
                        run.add_text('Тел. ')
                        run.add_text('(04848) 58142')
                        run.add_break()
                        run.add_text('Адреса електронної пошти: ')
                        # run.add_break()
                        add_hyperlink(paragraph3, "cnap_uspenovka@ukr.net", "https://cnap_uspenovka.ukr.net")
                        # run.add_text('cnap_uspenovka@ukr.net')
                        run2 = paragraph3.add_run()
                        run2.add_text('веб-сайт: ')
                        # run2.add_break()
                        run3 = paragraph3.add_run()
                        run3.add_text('https://uspenivska-gromada.gov.ua/')
                        run3.font.color.rgb = RGBColor(13, 109, 231)
                       
                        paragraph3.style = tmp_format
                        row3 = False
                        # print([paragraph.text, [paragraph.style.font.size]])
                    else:
                        pass
                        # print(paragraph.text)

                    if 'с. _________, вул. ______' in paragraph.text:
                        tmp_format = style2
                        # print(paragraph.runs[0].font.name )
                        paragraph.text = paragraph.text.replace('с. _________, вул. ______', 'с. Успенівка, вул. Кишинівська, 71')
                        # paragraph.style.element = tmp_format.element
                        paragraph.style = tmp_format

                        # print([paragraph.text])
        count += 1
        # print(count)
    return document
def parse_content(document):
    global out_
    count = 0
    flag, text = False, []
    for paragraph in document.paragraphs:
        # print([paragraph.text])
        if 'КАРТКА' in paragraph.text.upper():
            flag = True
        if 'Білгород-Дн' in paragraph.text.strip():
            flag = False
        if flag:
            if paragraph.text.strip() == '':
                count +=1
            if count == 2:
                flag = False
            # if paragraph.text.replace('\xa0\n', ' ').replace('  ', ' ').strip('\xa0\n\s') != '':
            text.append(paragraph.text.replace('\xa0\n', ' ').replace('  ', ' ').strip(' \xa0\n'))
    out_.append( ' '.join(text) )
    return document

# print(document.sections[0].iter_inner_content())
# Указываем путь к директории
# sub_cat = '/../../Цнап/Картки/ТК/'
sub_cat = ''
directory = os.getcwd() + f'{sub_cat}'.rstrip('/')

# Получаем список файлов
files = list( filter(lambda w: "~$" not in w, os.listdir(directory) ) )
# files.sort()
files.sort(key=natsort_keygen(alg=ns.REAL))
# print(files)
# Выводим список файлов
if 0:
    for name in files:
        # continue
        path = Path(f'{sub_cat}'+name)
        if path.suffix in ['.doc', '.rtf'] and '.docx' not in name:
            ok = ConvertRtfToDocx(directory, name)
            if ok:
                name = name.replace('doc', 'docx').replace('rtf', 'docx')
                path = Path(f'{sub_cat}'+name)
                out_ += name
                
        out_.append( str(path) )
        if path.suffix in {'.docx'}:
            d_ = open_doc(f'{sub_cat}'+name)
            new_d = parse_content(d_)
            new_d2 = parse_table(new_d)
            # print()
            new_d2.save('save/demo.docx')

if len(out_) > 0:
    with open(r"result.txt", "w", encoding="utf-8") as file:
        # print(out_)
        pattern = r'\s{1,}'
        file.write('Технологічні картки адміністративної послуги' + '\n' +
                   ('-' * 60 ) + '\n')
        for  line in out_:
            if '..\\..\\' in line:
                line = line.replace('..\\..\\', '\n')
                p = '01418'
            else:
                line = line.replace('Технологічна картка адміністративної послуги', '').strip()
                
                line = re.sub(pattern, ' ', line)
            file.write(line + '\n')
