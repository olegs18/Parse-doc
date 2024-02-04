# pip install python-docx
# pip install pywin32
# pip install natsort
# pip install Pillow
# https://python-docx.readthedocs.io/en/latest/user/text.html
from docx import Document
from docx.shared import Inches

from pathlib import Path
from PIL import Image
import os
import win32com.client
import shutil
import sys
from natsort import ns, natsort_keygen
import re
import csv_find
out_ = []

def image_to_jpg(image_path):
    path = Path(image_path)
    # if path.suffix not in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
    jpg_image_path = f'{path.parent / path.stem}.jpg'
    Image.open(image_path).convert('RGB').save(jpg_image_path)
    return jpg_image_path
    # return image_path

def ConvertRtfToDocx(rootDir, file):
    if os.path.exists(rootDir + "\\{f_}".format(f_= file.replace('doc', 'docx').replace('rtf', 'docx'))):
        if os.path.exists(rootDir + "\\" + file):
            shutil.move(rootDir + "\\" + file , rootDir + "\\arch\\" + file)
        return True
    word = win32com.client.Dispatch("Word.Application")
    wdFormatDocumentDefault = 16
    wdHeaderFooterPrimary = 1
    doc = word.Documents.Open(rootDir + "\\" + file)
    for pic in doc.InlineShapes:
        pic.LinkFormat.SavePictureWithDocument = True
    for hPic in doc.sections(1).headers(wdHeaderFooterPrimary).Range.InlineShapes:
        hPic.LinkFormat.SavePictureWithDocument = True
    doc.SaveAs(str(rootDir + "\\{f_}".format(f_= file.replace('doc', 'docx').replace('rtf', 'docx'))), FileFormat=wdFormatDocumentDefault)
    doc.Close()
    word.Quit()
    shutil.move(rootDir + "\\" + file , rootDir + "\\arch\\" + file)
    return True

def open_doc(file):
    doc = open(file,"rb")
    return Document(doc)
def parse_table(document):
    for table in document.tables:
        count = 0
        for rows in table.rows:
            _cell = ()
            for cell in rows.cells:
                if _cell == cell:
                    continue
                _cell = cell
                # cell.text += '__test__'
                for paragraph in cell.paragraphs:
                    for link in paragraph.hyperlinks:
                        pass
                            
                    print(paragraph.text)
        count += 1
        print(count)
def parse_content(document):
    global out_
    count = 0
    flag, text = False, []
    for paragraph in document.paragraphs:
        next_step = False
        # print([paragraph.text])
        if 'КАРТКА' in paragraph.text.upper():
            flag = True
        if 'Білгород-Дн' in paragraph.text.strip():
            if 'Білгород-Дн' not in re.sub(r'^Управління соціального захисту(.*)$', '',paragraph.text.strip()):
                flag = False                
            else:
                next_step = True
                paragraph.text = re.sub(r'Управління соціального захисту(.*)$', '', paragraph.text.strip())
        if flag:
            if next_step:
                flag = False
            if paragraph.text.strip() == '':
                count +=1
            if count == 2:
                flag = False
            # if paragraph.text.replace('\xa0\n', ' ').replace('  ', ' ').strip('\xa0\n\s') != '':
            text.append(paragraph.text.replace('\xa0\n', ' ').replace('  ', ' ').strip(' \xa0\n'))
    out_.append( ' '.join(text) )
    return document

# print(document.sections[0].iter_inner_content())
# Указываем путь к директории
sub_cat = '../../Цнап/Картки/ТК'
directory = os.getcwd() + f'/{sub_cat}'

# Получаем список файлов
files = list( filter(lambda w: "~$" not in w, os.listdir(directory) ) )
# files.sort()
files.sort(key=natsort_keygen(alg=ns.REAL))
# print(files)
# Выводим список файлов
for name in files:
    # continue
    if not ('.doc' in name or '.rtf' in name):
        continue
    path = Path(f'{sub_cat}/'+name)
    if path.suffix in ['.doc', '.rtf'] and '.docx' not in name:
        ok = ConvertRtfToDocx(directory, name)
        if ok:
            name = name.replace('doc', 'docx').replace('rtf', 'docx')
            path = Path(f'{sub_cat}/'+name)
            out_ += name
            
    out_.append( str(path) )
    if path.suffix in {'.docx'}:
        d_ = open_doc(f'{sub_cat}/'+name)
        new_d = parse_content(d_)
        # print()
        # new_d.save('save/demo.docx')
# parse_table()
if len(out_) > 0:
    c = j = t_= no_found = 0
    no_f = []
    with open(r"result.txt", "w", encoding="utf-8") as file:
        # print(out_)
        pattern = r'\s{1,}'
        file.write('Технологічні картки адміністративної послуги' + '\n' +
                   ('-' * 60 ) + '\n')
        for  line in out_:
            if '..\\..\\' in line:
                line = line.replace('..\\..\\', '\n')
                t = int( line.replace('Цнап\\Картки\\ТК\\ТК ', '').replace('.docx', '') )
                if t - t_ != 1:
                    no_f.append( str(int ( (t + t_)/2 )) )
                t_ = t
            else:
                line = line.replace('Технологічна картка адміністративної послуги', '').strip()
                line = re.sub(r'ннн', 'нн', line, flags=re.IGNORECASE)
                line = re.sub(r'посвідченЯ', 'посвідчення', line, flags=re.IGNORECASE)
                line = re.sub(pattern, ' ', line)
                usluga = csv_find.find_posl(line)
                
                if len(usluga) > 0:
                    if usluga[0] != '523p':
                        c += 1
                    else:
                        usluga[1] += ' [{}] '.format(usluga[3])
                        j += 1
                    line =  '{} '.format(usluga[1]) + line
            file.write(line + '\n')
if len(no_f) > 0:
    print(f'Нет {len(no_f)} карточек, а именно: ' + ', '.join(no_f))
print(c, j)