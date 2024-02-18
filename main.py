# pip install pysimplegui

# PyInstaller

import PySimpleGUI as sg
import main_ik
import main_tk
import analizze
from docx.shared import Pt
import subprocess
import time
import io
from contextlib import redirect_stdout
sg.theme("DarkBlue3")
radio_checked = b"iVBORw0KGgoAAAANSUhEUgAAABkAAAAZCAYAAADE6YVjAAAF40lEQVR4nI2Wf2yWVxXHv+fe+7y/3xbYWvpzhbGRCOkMLoRsjr21A2dI2BalTeaYxsyQ6GT+YTQuQRsy4zRGtmg2gzGNf+jinoK6sY2ZbNK3JQuSuWmiWx3ggBQKfTta+v58nueee/zjfQusMPD88yT3ued87sk593sPcCMTUblDYgZ80R9b90XnDomBiLphjOsEp8WBNQEiohUt2uuLhsji1Ut2zR8Dvq9HBgcZAPqPzK+ZD81DxWpwt2XucYIURCqa6FQmHnuryeBPY31N79dhvkbD77qQAV/0yCBx7tBMV0knn5oPooczyVR8Rcyi0zAS5FBhYDLQ+DDUKJWrtaxRf0hF87uObL3lzIL/J0IWNmx8c7Z/zsR/b7Rp25qex7aOuL09ayhhiECAs4xSyPLBxVD2T4bmQLkZURRNZaLi9nce7P4rfNG4AnQZIqJA5O4Zu5Cbk+TrHVRL/Hi1ie5cnjBgosAyWAAnAnEOEIcYCbRjOXy+an94XHlTHK8tcZUvvP1AR34h3mXIUL1DNm2eaTsXxN5t96R1uNdw15KkrgQMAqAgEAAiAuccnHOI2MFah4wWHJ+t8OMTWp8L9fn2uKwbP9JyHgCwm5wCgIG1IOwmdyH0no4lkq0/uQ22qzmhyzWGIUARINfqEBF4GrBaY83NKb2rJ7Amnlg+U+GnsZvcwNoRqmfSSOu+sYurT1Xdv7a3Oj10R5bKoZAhwAlAtBBTLmViLcMoQhBZfH84j7vXduLhDT3yvX+U5Y8fJXlVMlo7trX7GIZEqdwoFADMMn0pm057X2w3zjkQpH76mFFwTi4BRASWHYxWYCfY+dwb+M3L7+Bn/lHMViN6YDlcOpnwpgO1DQByfVAqXxgRACgHduMKz2JVxlBgHTxNIABnZopIJQwsuwaAYTTBOYcdzx7Ei2MT6O5Yih999bOA1rglAer2IpQZ9wBAvjAiCoODLCJkWXo6TIS4EoqsAwB899dv4q4nfouxf55GNh1HLYhgVD2zHc++jn2HP0D7sjR++c1+3PfpbhSrIZIa1KZCWJYVIkIYHOQF3dFOJJWAA4mAnQOzxdRHRZwtFPGVn76MN94+gZuWphBGFjueOYiR8f+gY1kGzz++CZ+7owuFi5X6nRBBHAxxkhodhQYA04AwQSoVJkTMcE7BMjD8nS0gIuwbn8BjP38Nz+3cjJH8BF7MT6Dz5gye37kJud5OFObKUASwc4gco+o8CFDp6wPXIb6viYhXv3rh5GSkP1UKQ1EaCEJG3NPY++374UTw0lvH8PU9B1GuRWi/KYNffWsz+no7MT1XgSLUa+YcSiHLmcgTD+FJIhL4vla5lgECgFQM4ycDQ8fmI/EgcCKoBhEIgr1PfB4P3nUbpueqaE7HsbeRwfRcGYoEzK7eEMI4XmSZjGKU8PQYAORaBsjkR+EAoNmofadL5d37zrLpbYoktEQeESq1EDFP4xff6Ec26WHL+pVXANAAOITWIUaRvFrQqlyphh0x3g8A+VE4ulIYe18pDLtE+mt72gt2Q0vCzIYCTwHOCYgIqbhBEFlUamG9kA15qVlGRjkcLQR21/kuo2rl4ROPdD+GAV9jZJA/pl259dOtU2LebTW27Zlbq7yyKabnQqnfTAiY619qACzX9SujGP+9GPCTp5bogjXnsiZc996/V0wvaNdVKvyZA2c2zqv0X1pRSz7ZVYnWL9UmFKKABdbVayUigGMYOChn5egM2z3nmr2CJCtZW73/vUd6Dl+twgvWeAfW/fn0vSXd9DttdHe/nsaWFmdXJkEJJUQQROxQDllOlEVeK2gzatvAbE+ng+L29x9dNf7J70nDFupz5/6T7dVY9qli6L6ciMWSXSZAOwWIE6PKhLM2jknroVwNqxmPXlgSXPjB3x9dM7UYcE1IPaPLb/WGA9O3zzM9VAr5XhvZlQ6SIaGSUfRh0jP5ZRS+9Ldt3ccW+/1/JkJYNK0oAg6JmKtmIN+/7rRyYxuqz12LgfD9+tw1dOO563+8H1VJkK2keQAAAABJRU5ErkJggg=="
radio_unchecked = b"iVBORw0KGgoAAAANSUhEUgAAABkAAAAZCAYAAADE6YVjAAAEwElEQVR4nI1W3W9URRT/nZm7ZXdpbajdWpCAjcFEqw88+CACrgaBmFBIwI3fPPpPaJYND/wjYsxFYgwP+BV2kY9gNCIJIhEIBZSWLl3aprvde2fOOT7c3W27fNSTTO7cMzO/35wz55wZYAVRVVMuaxCGoV2qD8PQlsvlQFXNShhPAqduYEr0lrrmhmFoVbVbvWzdQxKGoS0UCgwAFy6PvySx27cQRVvY80YGZyHaIKJbPUHqvCF8k3/tlb+61z2RJAzVFgrE5QuX1q9K9x6Oouj9TCazKmUBawiAglkQO0bsPOqNejOw9qsoan62Z8eWfx9FRMsJkgnnfrv6FgXBUWOD4UzAWJsb8L3ZNFlrCQSwZ8TO6excXe/eux/UY0EcuQkXRx/t3fX6qW6iDomqGiKS87///QaM/Q7K6efXD7rBgf5AVcl7hgBQEYgqVAQEgqroZLXmb9yeTLGgKRztHtu5/XQbr0NSVDU4dAhvj703LGouBpaGXhwZ5v6nem0cO2gCB002AxGBiICZwSwIrEVtZpav3LhjneN76YxsvnDq1D0AKJVKYgBg9NgxKpVKIkpH0ulVQyPrBvxTfb02ih2ICESAdp2darJHIkIUx+jrXW03rB30PT09zzTm5UipVJLR0VECAGqb9csfV16oN3H56f60Hd20gZzzRJR4UzvAusySxBoBi8A5DyLolWvjOv1gjldnUqN7duavFYtFYyoVGACIvd2fzWZSw4P9IqKkLfBugu4GKFSSr4hSbqBfMplMaiFyBwAgn88bU60eUwCI43hbYIBsJk2e+bHAiQVL/xWiSTB4ZmQzabKG4B1vBYBqtapBoVBgVaUfz13aaI3CEBGzgAjouEuXg3bARSG6pImADJEhwLN/TlWJiDhoecOqSHYpUIJPHYclY4CqdBElZ6Otfse9otlKBRaAb5OwqjbaYSnatqKzpEXQAleFsIAlCWERBbfyR4TBwlDVRj4PBgAThqElIgVhPPaicew02R0vi6ClESWcALEkkbV0bhQ7dZ4VpONEpGEYWpPL5QgArLVnYsc0N99QAuC5nWy8JPEYvtW4PS6LfVXFfL2hznkyxv4MALlcjkwlnxcACCj4ul6fjyeqNeOZ1Xu/COoXwX0XkbDAs8B7BjPrVLVm6vVGDOXjAFCpVMSUiCQMQ/vmlpevE+nRyJOZul9jYwix84sEfrG1d94h9A5EQHW6xrEXYwhffFLYe/3dMLSlUkmS2lUsGgB4Nf/OEIleJEPDI88Ocl/vauu8b5UQdA69nS/t2mWIMDM3x+P/TFp2flKM3Tz+569T7dr1UBU+8dPZbWRS30M4s25ojVvT3xcIlNpRpCpd+cI6XZvxd6emUyrUEPW7DhbGzi6twp37mVpu27Nj65lmo7lbgDsT9+dSV2/cotqDWR/HMYt4ERHx7CWKIq7NzPrrN2/TVG0uBcVt56PdBwtjZ1sRKx3sruLaubiOnzy51tq+wy6KP0j19GSsAQwtlnrPjNgxmgvNBWvNl41m8/NPP94/seLN2E0EACd+qGxyse5runi7Zz+iLL2imLcGN1PWnhYNvv3wwM5r3ev+lzzqtdLSB926lV4rK0qxWDTlcvmx7652ZD5J/gNoDCDS80MCGwAAAABJRU5ErkJggg=="


layout = [
    [
        sg.Text("Работа с карточками", (20, 1), font="Arial 15 italic"),
    ],
    [sg.Text("Выбирите тип карточек", font="arial 12 normal")],
    [
        sg.Radio(
            "Інформаційні",
            "type_data",
            key="type",
            font="arial 12 normal",
            tooltip="Если будете работать с інфомационными карточками",
        ),
        sg.Radio(
            "Технологічні",
            "type_data",
            key="type",
            font="arial 12 normal",
            tooltip="Если будете работать с технологическими карточками",
        ),
    ],
    [
        sg.Checkbox(
            "Редактировать карточки под ЦНАП (отредактированные карточки будут в папке Edited)",
            key="edited",
        )
    ],
    [
        sg.Text("Выбирите пук к папке", font="arial 12 normal"),
        sg.FolderBrowse(
            "Обзор",
            font="arial 12 normal",
            tooltip="Укажите путь к файлам, где расположены карточки, с которыми необходимо работать",
        ),
    ],
    # [sg.HorizontalSeparator()],
    [sg.Output(size=(88, 20), key='-OUTPUT-')],
    [sg.Button("Обработать"), sg.Cancel("Отмена")],
    [sg.HorizontalSeparator()],
    [
        sg.Text("Работа с услугами", (20, 1), font="Arial 15 italic"),
    ],
    [
        [
            sg.InputText(key="usluga"),
            sg.Button(
                "Получить сведения", k="-check1-"
            ),  # ,sg.Image(radio_checked, enable_events=True, k='-check-', metadata=True),
        ]
    ],
    [
        sg.Text("Работа с услугами (массово)", (30, 1), font="Arial 15 italic"),
    ],
    [
        [
            sg.Multiline(default_text='', key='input_text', size=(40, 10), pad=((0, 0), (0, 5)) ),
            sg.Button(
                "Получить сведения", k="-check2-"
            )
        ]
    ],
]


def get_sub(values):
    if "\\" in main_ik.os.getcwd():
        sp_path = main_ik.os.getcwd().split("\\")
    else:
        sp_path = main_ik.os.getcwd().split("/")
    count = 0
    for l in sp_path[::-1]:
        if l not in values["Обзор"]:
            count += 1
        else:
            break
    if count:
        for i in range(1, count + 1):
            sp_path[-i] = ".."
        pre_sub = values["Обзор"].split(sp_path[-(count + 1)])
        if len(pre_sub) == 2:
            sub = "/".join([".." for _ in range(count)]) + pre_sub[1]
        else:
            sub = ""
    else:
        sub = ""
    # print(sub)
    return sub


import textwrap

width = 40


def wrap(string, lenght=width):
    return "\n".join(textwrap.wrap(string, lenght))
def copy_text(element):
    element.Widget.event_generate('<Control-c>')

def select_all_text(element):
    element.Widget.tag_add('sel', '1.0', 'end')

def popup_text(filename, text, sp):
    if filename == None:
        filename = "Буфер"
    if isinstance(text, list):
        # text = list(map(lambda x: list(map(wrap, x)), text))
        if len(sp) > 0:
            tt = 'Услуга обязательна по 523р. '
            sp_ = sp[0]
            ss = [ sg.Text(tt, font='"Times New Roman" 12'), sg.Text( sp_[0] , font='"Times New Roman" 12 bold'), sg.Text( "\n".join(textwrap.wrap(sp_[1], 80)), font='"Times New Roman" 12' ), sg.Text(sp_[2], font='"Times New Roman" 12') ]
        else:
            tt = ''
            sp_ = sp
            ss = [sg.Text('Услуга не обязательна (по 523р)', font='"Times New Roman" 12')]
        layout = [
            ss,
            [
                sg.Table(
                    values=text,
                    headings=[str(i + 1) for i in range(len(text[0]))],
                    max_col_width=60,
                    background_color="black",
                    auto_size_columns=True,
                    display_row_numbers=True,
                    justification="l",
                    num_rows=len(text),
                    # alternating_row_color="black",
                    key="-TABLE-",
                    expand_x=False,
                    expand_y=True,
                    # enable_events=True,
                    enable_click_events=True,
                    # row_height=35=
                )
            ],
            [sg.Button("Скопировать всю таблицу")],
            [
                sg.Multiline(
                    size=(160, 10), key="-FULLTEXT-", visible=False, autoscroll=True,
                     right_click_menu=['&Right', ['&Copy', 'Select &All']],
                     enable_events=True
                )
            ],
        ]
    else:
        layout = [
            [sg.Text(filename, font='"Times New Roman" 12')],
            [sg.Multiline(text, size=(180, 25))],
        ]

    # layout.append([sg.Multiline(text, size=(80, 25))])
    location = sg.Window.get_screen_size()
    win = sg.Window(
        f"Результат обработки",
        layout,
        modal=True,
        finalize=True,
        titlebar_icon=radio_checked,
        resizable=True,
        ttk_theme="clam",
        location=location,
    )

    # # Update the GUI for the Table
    # style_name = table.table_ttk_style_name + '.Heading'
    # style = table.ttk_style
    # foreground = style.configure(style_name)['foreground']
    # style.configure(style_name, foreground=foreground)

    # Relocate window to center of the screen
    win.refresh()
    win.move_to_center()

    while True:
        event, values = win.read()
        if event == sg.WINDOW_CLOSED:
            break
        elif '-TABLE-' in event and '+CLICKED+' in event:
            if event[2][0] != None and event[2][1] != None:
                full_text = text[event[2][0]][event[2][1]]
                win["-FULLTEXT-"].update(value=full_text)
                win["-FULLTEXT-"].update(visible=True)
        elif event == "Скопировать всю таблицу":
            clipboard_data = ''
            for row in text:
                clipboard_data += '\t'.join(map(str, row)) + '\n'
            sg.clipboard_set(clipboard_data)
        elif event == '-FULLTEXT-Right':
            right_click_event = values[event]
            if right_click_event == 'Copy':
                copy_text(win['-FULLTEXT-'])
            elif right_click_event == 'Select All':
                select_all_text(win['-FULLTEXT-'])
        if event == 'Copy':
            copy_text(win['-FULLTEXT-'])
        elif event == 'Select All':
            select_all_text(win['-FULLTEXT-'])
        # print(event)
    win.close()


f_n = radio_checked
window = sg.Window("Обработка карточек", layout, resizable=True, titlebar_icon=f_n)
while True:  # The Event Loop
    event, values = window.read()
    # print(event, values) #debug
    # print(event)
    if event in ("-check-", "-check1-"):
        if not values["usluga"]:
            arg = sg.popup("Услуга не указана. Укажите услугу")
        else:
            # window[0].update(text_color=sg.theme_text_element_background_color(), background_color=sg.theme_text_color())
            g = analizze.get_data(values["usluga"])
            
            data_523 = analizze.get_all_data_523()
            txt = []
            sp_ = [word for word in data_523 if values["usluga"] in word[0]]
            if len(g) > 0:
                for text in g:
                    sp = list(text)
                    txt.append([sp[0], text[sp[0]]])
            else:
                if len(sp_) > 0:
                    txt.append([sp_[0][0], sp_[0][1] + ' ' + sp_[0][2]])
            try:
                # with open(filename, "rt", encoding="utf-8") as f:
                #     text = f.read()
                popup_text(None, txt, sp_)
            except Exception as e:
                print("Error: ", e)
        pass

    if event is not None and event in "Обработать":
        # print(values, values['type'], values['type0'], values['Обзор'])
        if values["type"]:
            if "ІК" not in values["Обзор"]:
                # arg = sg.popup('Вы уверены? В пути не встречается папка ІК')
                arg = "OK"
                if arg.upper() == "OK":
                    # print(arg)
                    print("Обрабатываем Інформаційні картки")

                    if values["Обзор"] != "":
                        # pass
                        sub = get_sub(values)
                        main_ik.pr_s(values["Обзор"], sub)
                    # main_ik.pr_d()
                    main_ik.go(values["edited"])
                    # if len(layout) == 6:
                    #     layout.append([])
                    #     layout[5] , layout[6] = layout[6] , layout[5]
                    filename = main_ik.os.getcwd() + "/result.txt"
                    if main_ik.Path(filename).is_file():
                        try:
                            with open(filename, "rt", encoding="utf-8") as f:
                                text = f.read()
                            popup_text(filename, text, [])
                        except Exception as e:
                            print("Error: ", e)
            else:
                print("Обрабатываем Інформаційні картки")
                if values["Обзор"]:
                    sub = get_sub(values)
                    main_ik.pr_s(values["Обзор"], sub)
                # main_ik.pr_d()
                main_ik.go(values["edited"])
                filename = main_ik.os.getcwd() + "/result.txt"
                if main_ik.Path(filename).is_file():
                    try:
                        with open(filename, "rt", encoding="utf-8") as f:
                            text = f.read()
                        popup_text(filename, text, [])
                    except Exception as e:
                        print("Error: ", e)
        elif values["type0"]:
            if "ТК" not in values["Обзор"]:
                # arg = sg.popup('Вы уверены? В пути не встречается папка ТК')
                arg = "OK"

                if arg.upper() == "OK":
                    # print(arg)
                    print("Обрабатываем Технологічні картки")
                    if values["Обзор"]:
                        sub = get_sub(values)
                        main_tk.pr_s(values["Обзор"], sub)
                    main_tk.go(values["edited"])
                    filename = main_ik.os.getcwd() + "/result.txt"
                    if main_ik.Path(filename).is_file():
                        try:
                            with open(filename, "rt", encoding="utf-8") as f:
                                text = f.read()
                            popup_text(filename, text, [])
                        except Exception as e:
                            print("Error: ", e)

            else:
                print("Обрабатываем Технологічні картки")
                if values["Обзор"]:
                    sub = get_sub(values)
                    main_tk.pr_s(values["Обзор"], sub)
                main_tk.go(values["edited"])
                filename = main_ik.os.getcwd() + "/result.txt"
                if main_ik.Path(filename).is_file():
                    try:
                        with open(filename, "rt", encoding="utf-8") as f:
                            text = f.read()
                        popup_text(filename, text, [])
                    except Exception as e:
                        print("Error: ", e)
    elif event in (sg.WIN_CLOSED, "Отмена", "Cancel"):
        break
    elif event == '-check2-':
        posluga = []
        if not values['input_text']:
            arg = sg.popup("Услуга не указана. Укажите услугу")
        else:
            ls_input_text = values['input_text'].split('\n')
            for input_text in ls_input_text:
                if input_text.strip() == '':
                    continue
                g = analizze.get_data(input_text)
                data_523 = analizze.get_all_data_523()
                txt = []
                sp_ = [word for word in data_523 if input_text in word[0]]
                if len(g) > 0:
                    for text in g:
                        sp = list(text)
                        txt.append([sp[0], text[sp[0]]])
                else:
                    if len(sp_) > 0:
                        txt.append([sp_[0][0], sp_[0][1] + ' ' + sp_[0][2]])
                try:
                    # with open(filename, "rt", encoding="utf-8") as f:
                    #     text = f.read()
                    with io.StringIO() as buf, redirect_stdout(buf):
                        print(f'Услуга {input_text} обработана')
                        # Получаем вывод и обновляем элемент sg.Output
                        output_text = buf.getvalue()
                        window['-OUTPUT-'].update(output_text, append=True)
                        window.refresh()

                    posluga.append( {input_text: [
                        txt, sp_
                    ]})
                except Exception as e:
                    with io.StringIO() as buf, redirect_stdout(buf):
                        print("Error: ", e)
                        # Получаем вывод и обновляем элемент sg.Output
                        output_text = buf.getvalue()
                        window['-OUTPUT-'].update(output_text, append=True)
                        window.refresh()
                    
            document = main_ik.Document()
            # Добавляем раздел (section) к документу
            section = document.sections[0]

            # Устанавливаем отступы для раздела (в точках)
            section.left_margin = Pt(30)  # Левое поле
            section.right_margin = Pt(15)  # Правое поле
            section.top_margin = Pt(15)  # Верхнее поле
            section.bottom_margin = Pt(15)  # Нижнее поле
            
            document.add_heading('Послуги', 0)
            with io.StringIO() as buf, redirect_stdout(buf):
                print('Приступаем к формированию Word документа')
                # Получаем вывод и обновляем элемент sg.Output
                output_text = buf.getvalue()
                window['-OUTPUT-'].update(output_text, append=True)
                window.refresh()
            count = 0
            if len(posluga) < 10:
                pr_del = 5
            elif len(posluga) >= 10 and len(posluga) < 20:
                pr_del = 10
            elif len(posluga) >= 20 and len(posluga) < 40:
                pr_del = 15
            elif len(posluga) >= 40 and len(posluga) < 80:
                pr_del = 20
            elif len(posluga) >= 80:
                pr_del = 25
            for doc in posluga:
                count +=1
                if int((count / len(posluga) ) * 100 ) % pr_del == 0:
                    with io.StringIO() as buf, redirect_stdout(buf):
                        print(f'Формирование обработано на {str(int((count / len(posluga) ) * 100 ) )}%')
                        # Получаем вывод и обновляем элемент sg.Output
                        output_text = buf.getvalue()
                        window['-OUTPUT-'].update(output_text, append=True)
                        window.refresh()

                usl = list(doc)
                usls= doc[usl[0]]
                usl_523 = usls[1]
                usl_table = usls[0]
                document.add_heading(f'Послуга {usl[0]}', level=1)
                p = document.add_paragraph('Услуга ')
                if len(usl_523) > 0 and usl[0] in usl_523[0]:
                    s = usl_523[0].pop(0)
                    p.add_run(s).bold = True
                    tt = f' обязательна по 523р. \n\r' + '. '.join(usl_523[0])
                else:
                    p.add_run(usl[0]).bold = True
                    tt = f' не обязательна по 523р. '
                p.add_run(tt).normal = True

                table = document.add_table(rows=1, cols=2, style='Colorful List')
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Назва'
                hdr_cells[1].text = 'Значення'
                for qty, id in usl_table:
                    row_cells = table.add_row().cells
                    # row_cells[0].text = str(qty)
                    p = row_cells[0].paragraphs[0]
                    p.add_run(str(qty)).bold = True
                    row_cells[0].width = main_ik.Inches(2.0)
                    row_cells[1].text = id
                    row_cells[1].width = main_ik.Inches(6.0)
                if len(posluga) > 1 and posluga[len(posluga) - 1] != usl[0]:
                    document.add_page_break()
                    # time.sleep(1)
            try:
                document.save('info.docx')
            except Exception as e:
                print(f"Ошибка при сохранении информации: {e}")    
            try:
                file_path = 'info.docx'
                subprocess.Popen(["start", " ", file_path], shell=True)
            except Exception as e:
                print(f"Ошибка при открытии файла: {e}")


window.close()
